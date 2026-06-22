from docx import Document
from docx.oxml.ns import qn
from typing import List
from .base_parser import BaseParser, ParsedChunk


class DOCXParser(BaseParser):
    def parse(self, file_path: str) -> List[ParsedChunk]:
        chunks = []
        doc = Document(file_path)
        chunks.extend(self._extract_paragraphs(doc, file_path))
        chunks.extend(self._extract_tables(doc, file_path))
        return chunks

    def _extract_paragraphs(self, doc, file_path: str) -> List[ParsedChunk]:
        chunks = []
        buffer = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style = para.style.name.lower()
            # Flush buffer on headings
            if "heading" in style and buffer:
                meta = self._base_metadata(file_path, "docx")
                meta["chunk_type"] = "prose"
                chunks.append(ParsedChunk(text=" ".join(buffer), metadata=meta))
                buffer = []
            buffer.append(text)

        if buffer:
            meta = self._base_metadata(file_path, "docx")
            meta["chunk_type"] = "prose"
            chunks.append(ParsedChunk(text=" ".join(buffer), metadata=meta))

        return chunks

    def _extract_tables(self, doc, file_path: str) -> List[ParsedChunk]:
        chunks = []
        for table_idx, table in enumerate(doc.tables):
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if len(rows) < 2:
                continue
            headers = rows[0]
            for row_idx, row in enumerate(rows[1:], start=1):
                row_text = " | ".join(
                    f"{headers[i]}={val}"
                    for i, val in enumerate(row)
                    if val
                )
                if not row_text.strip():
                    continue
                meta = self._base_metadata(file_path, "docx")
                meta.update({
                    "table_index": table_idx,
                    "row_index": row_idx,
                    "headers": headers,
                    "chunk_type": "table_row",
                })
                chunks.append(ParsedChunk(text=row_text, metadata=meta))
        return chunks